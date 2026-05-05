# -*- coding: utf-8 -*-
from typing import Optional
from docx import Document

from .base_reader import BaseReader
from util.logging.logger import get_logger

logger = get_logger(__name__)

class DocxReader(BaseReader):
    """Word文档读取器类"""
    
    def read_text(self, file_path: str) -> Optional[str]:
        """
        读取Word文档中的纯文本内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            str: Word文档中的文本内容，如果读取失败则返回None
        """
        try:
            if not self._check_file_exists(file_path):
                return None
                
            # 打开Word文档
            doc = Document(file_path)
            text_content = []
            
            # 读取所有段落
            for para in doc.paragraphs:
                if para.text.strip():  # 只添加非空段落
                    text_content.append(para.text)
            
            # 读取所有表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # 只添加非空单元格
                            row_text.append(cell.text.strip())
                    if row_text:  # 只添加非空行
                        text_content.append(' | '.join(row_text))
            
            # 合并所有文本内容
            full_text = '\n'.join(text_content)
            logger.info(f"成功读取Word文档文本内容: {file_path}")
            return full_text
                
        except Exception as e:
            logger.error(f"读取Word文档文本内容失败: {str(e)}")
            return None

# 创建全局单例实例
docx_reader = DocxReader() 