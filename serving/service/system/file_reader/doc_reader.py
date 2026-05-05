# -*- coding: utf-8 -*-
from typing import Optional
import os
import platform
import subprocess
import tempfile
from pathlib import Path

from .base_reader import BaseReader
from util.logging.logger import get_logger

logger = get_logger(__name__)


class DocReader(BaseReader):
    """旧版Word文档(.doc)读取器类"""

    def _check_system_dependencies(self) -> bool:
        """
        检查系统依赖是否满足

        Returns:
            bool: 系统依赖是否满足
        """
        system = platform.system().lower()
        if system == 'windows':
            try:
                import win32com.client
                import pythoncom
            except ImportError:
                logger.error("Windows系统下需要安装pywin32: pip install pywin32")
                return False
        # Linux: LibreOffice 可用性由 _read_text_linux 内部检测并报告具体错误
        return True

    def read_text(self, file_path: str) -> Optional[str]:
        """
        读取旧版Word文档(.doc)中的纯文本内容

        Args:
            file_path: Word文档路径

        Returns:
            str: Word文档中的文本内容

        Raises:
            RuntimeError: Linux环境下LibreOffice不可用或转换失败时
        """
        if not self._check_file_exists(file_path):
            return None

        if not self._check_system_dependencies():
            return None

        system = platform.system().lower()
        if system == 'windows':
            return self._read_text_windows(file_path)
        elif system == 'linux':
            # 不再静默返回 None，让异常向上传播，调用方可感知并处理
            return self._read_text_linux(file_path)
        else:
            raise RuntimeError(
                f"不支持的操作系统: {system}，无法读取.doc文件: {file_path}"
            )

    # ==================== Linux 实现 ====================

    def _read_text_linux(self, file_path: str) -> str:
        """
        Linux系统下通过 LibreOffice headless 将 .doc 转换为 .docx 后读取文本

        转换在系统临时目录中进行，转换完成后自动清理临时文件。

        Args:
            file_path: .doc 文件路径

        Returns:
            str: 提取的纯文本内容

        Raises:
            RuntimeError: LibreOffice不可用、转换失败或超时时
        """
        abs_path = Path(file_path).resolve()

        # 检查 LibreOffice 是否可用
        libreoffice_cmd = self._find_libreoffice()
        if not libreoffice_cmd:
            raise RuntimeError(
                f"Linux环境下无法读取.doc文件: {abs_path}，"
                "请安装LibreOffice（apt-get install -y libreoffice）或将文件转为.docx格式"
            )

        # 在临时目录中执行转换，退出 with 块时自动清理
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            try:
                result = subprocess.run(
                    [
                        libreoffice_cmd, "--headless",
                        "--convert-to", "docx",
                        "--outdir", str(tmp_path),
                        str(abs_path)
                    ],
                    timeout=60,
                    capture_output=True,
                    text=True
                )
            except FileNotFoundError:
                raise RuntimeError(
                    f"Linux环境下无法读取.doc文件: {abs_path}，"
                    "请安装LibreOffice（apt-get install -y libreoffice）或将文件转为.docx格式"
                )
            except subprocess.TimeoutExpired:
                raise RuntimeError(
                    f"LibreOffice转换超时（>60s）: {abs_path}，文件可能过大或已损坏"
                )

            if result.returncode != 0:
                raise RuntimeError(
                    f"LibreOffice转换失败(returncode={result.returncode}): "
                    f"file={abs_path}, stderr={result.stderr.strip()}"
                )

            # LibreOffice 输出文件名与原文件同名，扩展名改为 .docx
            docx_path = tmp_path / (abs_path.stem + ".docx")
            if not docx_path.exists():
                raise RuntimeError(
                    f"LibreOffice转换后未找到输出文件: {docx_path}，"
                    f"stdout={result.stdout.strip()}"
                )

            logger.info(f"LibreOffice转换成功: {abs_path} → {docx_path}")
            return self._extract_text_from_docx(str(docx_path))

    def _find_libreoffice(self) -> Optional[str]:
        """
        查找 LibreOffice 可执行文件

        依次尝试 'libreoffice' 和 'soffice'（两者均为常见命令名）

        Returns:
            str: 可用的命令名，找不到则返回 None
        """
        for cmd in ["libreoffice", "soffice"]:
            try:
                result = subprocess.run(
                    ["which", cmd],
                    capture_output=True,
                    text=True,
                    timeout=5
                )
                if result.returncode == 0 and result.stdout.strip():
                    logger.debug(f"找到LibreOffice命令: {cmd} → {result.stdout.strip()}")
                    return cmd
            except Exception:
                continue
        return None

    def _extract_text_from_docx(self, docx_path: str) -> str:
        """
        使用 python-docx 提取 .docx 文件的纯文本（段落 + 表格）

        复用与 DocxReader 相同的逻辑，保证输出格式一致。

        Args:
            docx_path: .docx 文件路径

        Returns:
            str: 提取的纯文本，段落间以换行分隔，表格行以 ' | ' 分隔
        """
        from docx import Document

        doc = Document(docx_path)
        text_content = []

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if row_text:
                    text_content.append(' | '.join(row_text))

        full_text = '\n'.join(text_content)
        logger.info(f"成功从转换后的.docx提取文本，共 {len(text_content)} 段")
        return full_text

    # ==================== Windows 实现（不改动）====================

    def _read_text_windows(self, file_path: str) -> Optional[str]:
        """
        Windows系统下读取.doc文件

        Args:
            file_path: Word文档路径

        Returns:
            str: Word文档中的文本内容，如果读取失败则返回None
        """
        try:
            import win32com.client
            import pythoncom

            abs_path = str(Path(file_path).resolve())

            if not os.path.exists(abs_path):
                logger.error(f"文件不存在: {abs_path}")
                return None

            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            try:
                doc = word.Documents.Open(abs_path)
                text_content = doc.Content.Text
                doc.Close()

                if text_content:
                    logger.info(f"成功读取.doc文档文本内容: {abs_path}")
                    return text_content.strip()
                else:
                    logger.warning(f"文档内容为空: {abs_path}")
                    return None

            finally:
                word.Quit()
                pythoncom.CoUninitialize()

        except Exception as e:
            logger.error(f"Windows系统下读取.doc文档失败: {str(e)}")
            return None


# 创建全局单例实例
doc_reader = DocReader()
